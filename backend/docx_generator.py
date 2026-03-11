"""
RERA Reports Word Document Generator - Goa State Official Format
Generates Word (.docx) versions of Form-1, Form-3, Form-4, and Annexure-A
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from io import BytesIO
from datetime import datetime


# ─── Helpers ────────────────────────────────────────────────────────────────

def format_indian_number(num):
    if num is None or num == 0:
        return "0"
    num = int(num)
    is_negative = num < 0
    num = abs(num)
    s = str(num)
    if len(s) <= 3:
        result = s
    else:
        result = s[-3:]
        s = s[:-3]
        while s:
            result = s[-2:] + "," + result
            s = s[:-2]
    return ("-" + result) if is_negative else result


def get_cat_avg(cat_data):
    if not cat_data:
        return 0
    vals = [v.get("completion", 0) for v in cat_data.values() if isinstance(v, dict)]
    return sum(vals) / len(vals) if vals else 0


# ─── Cell Styling Helpers ─────────────────────────────────────────────────────

def shade_cell(cell, hex_color):
    """Set background colour for a table cell."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def border_cell(cell, color="000000", size="4"):
    """Add thin borders to all sides of a cell."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in ("top", "bottom", "left", "right"):
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), size)
        b.set(qn("w:color"), color)
        tcBorders.append(b)
    tcPr.append(tcBorders)


def set_col_width(table, col_index, width_cm):
    """Set width of a specific column in a table."""
    from docx.oxml.ns import qn as _qn
    from docx.oxml import OxmlElement as _OE
    from docx.shared import Cm as _Cm
    for row in table.rows:
        cell = row.cells[col_index]
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcW = _OE("w:tcW")
        tcW.set(_qn("w:w"), str(int(_Cm(width_cm).emu / 914.4 * 20)))
        tcW.set(_qn("w:type"), "dxa")
        tcPr.append(tcW)


# ─── Paragraph / Run Helpers ──────────────────────────────────────────────────

def add_para(doc, text="", size=10, bold=False, italic=False,
             align=WD_ALIGN_PARAGRAPH.LEFT, color=None, space_after=Pt(4)):
    para = doc.add_paragraph()
    para.alignment = align
    para.paragraph_format.space_after = space_after
    run = para.add_run(text)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)
    return para


def add_title_block(doc, form_num, title_text, subtitle_text):
    for line in [
        "The Goa Real Estate (Regulation and Development)",
        "(Registration of Real Estate Projects, Registration of Real Estate Agents,",
        "Rates of Interest and Disclosures on Website) Rules 2017",
    ]:
        p = add_para(doc, line, size=9, align=WD_ALIGN_PARAGRAPH.CENTER)

    add_para(doc, "")
    add_para(doc, form_num, size=14, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc, "(See Rule 5 (1) (a) (ii))", size=9, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc, title_text, size=12, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc, subtitle_text, size=9, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc, "")


def add_info_table(doc, rows):
    """Rows: list of (label, value) tuples."""
    t = doc.add_table(rows=len(rows), cols=2)
    t.style = "Table Grid"
    for i, (lbl, val) in enumerate(rows):
        c0, c1 = t.cell(i, 0), t.cell(i, 1)
        c0.text = lbl
        c0.paragraphs[0].runs[0].font.bold = True
        c0.paragraphs[0].runs[0].font.size = Pt(10)
        c1.text = str(val) if val else ""
        c1.paragraphs[0].runs[0].font.size = Pt(10)
        shade_cell(c0, "D9E1F2")
    return t


def add_section_heading(doc, text, color=(0x1F, 0x38, 0x64)):
    p = add_para(doc, text, size=11, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, color=color)
    return p


def add_data_table(doc, rows, header_bg="4472C4", total_bg="E2EFDA"):
    """rows[0] = header. rows[-1] can be a total row (handled by caller)."""
    n_cols = len(rows[0])
    t = doc.add_table(rows=len(rows), cols=n_cols)
    t.style = "Table Grid"

    for r_idx, row_data in enumerate(rows):
        is_header = r_idx == 0
        is_total  = r_idx == len(rows) - 1 and len(rows) > 2

        for c_idx, val in enumerate(row_data):
            cell = t.cell(r_idx, c_idx)
            cell.text = str(val) if val is not None else ""
            para = cell.paragraphs[0]
            run  = para.runs[0] if para.runs else para.add_run(str(val) if val is not None else "")
            run.font.size = Pt(9)
            run.font.bold = is_header or is_total
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            border_cell(cell)

            if is_header:
                shade_cell(cell, header_bg)
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            elif is_total:
                shade_cell(cell, total_bg)

    return t


def section_margins(doc):
    for section in doc.sections:
        section.top_margin    = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin   = Cm(2.0)
        section.right_margin  = Cm(2.0)


# ─── FORM 1 ──────────────────────────────────────────────────────────────────

def generate_form1_docx(project, buildings, construction_progress, infrastructure_progress, quarter, year):
    doc = Document()
    section_margins(doc)

    add_title_block(doc, "FORM 1", "ARCHITECT'S / LICENSED SURVEYOR'S CERTIFICATE",
                    "(To be submitted at the time of Registration and for withdrawal of Money from Designated Account)")

    # "To" block
    add_para(doc, "To", bold=True)
    for line in [
        project.get("promoter_name", "________________________"),
        project.get("promoter_address", project.get("address", "________________________")),
        f"{project.get('village', '')}, {project.get('taluka', '')}, {project.get('district', 'Goa')}",
    ]:
        add_para(doc, line, size=10)

    p_date = doc.add_paragraph()
    p_date.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p_date.add_run(f"Date: {datetime.now().strftime('%d/%m/%Y')}")
    r.font.bold = True
    r.font.size = Pt(10)

    add_para(doc, "")

    # Project info table
    add_info_table(doc, [
        ("Project Name",            project.get("project_name", "")),
        ("RERA Registration No.",   project.get("rera_number", "")),
        ("Report Period",           f"{quarter} {year}"),
        ("Promoter Name",           project.get("promoter_name", "")),
        ("Architect Name",          project.get("architect_name", "")),
        ("Architect License No.",   project.get("architect_license", "")),
        ("Survey No. / Plot No.",   project.get("survey_number", "")),
        ("Total Area (sq.m.)",      str(project.get("total_area", ""))),
        ("Village / Taluka",        f"{project.get('village', '')}, {project.get('taluka', '')}"),
        ("District",                project.get("district", "North Goa")),
    ])

    add_para(doc, "")

    # Assignment text
    num_buildings = len(buildings) if buildings else "__"
    phase = project.get("phase", "1")
    add_para(doc, (
        f"I/We {project.get('architect_name', '________________________')} have undertaken assignment as "
        f"Architect / Licensed Surveyor of certifying Percentage of Completion of Construction Work of "
        f"{num_buildings} Building(s) / Wing(s) of the {phase} Phase of the Project, "
        f"RERA No. {project.get('rera_number', '________')}."
    ), size=10)

    # TABLE A for each building
    progress_lookup = {p.get("building_id"): p for p in (construction_progress or [])}

    for building in (buildings or []):
        bname    = building.get("building_name", "Building")
        progress = progress_lookup.get(building.get("building_id"), {})
        ta       = progress.get("tower_activities", {})

        plinth   = ta.get("plinth_completion", {})
        slab     = ta.get("slab_completion", {})
        brickwk  = ta.get("brickwork_plastering", {})
        plumbing = ta.get("plumbing", {})
        elec     = ta.get("electrical_works", {})
        windows  = ta.get("window_works", {})
        tiling   = ta.get("tiling_flooring", {})
        doors    = ta.get("door_shutter_fixing", {})
        waterp   = ta.get("water_proofing", {})
        painting = ta.get("painting", {})
        carpark  = ta.get("carpark", {})

        add_para(doc, "")
        add_section_heading(
            doc,
            f"TABLE A: Building / Wing — {bname}  (Overall Completion: {progress.get('overall_completion', 0):.1f}%)"
        )

        rows_a = [
            ["Sr. No.", "Tasks / Activity", "Percentage of Work Done"],
            ["1",  "Excavation",
             f"{plinth.get('excavation', {}).get('completion', 0):.0f}%"],
            ["2",  f"{building.get('basements', 0)} Basement(s) and Plinth",
             f"{get_cat_avg(plinth):.0f}%"],
            ["3",  f"{building.get('residential_floors', 0)} Slabs of Super Structure",
             f"{get_cat_avg(slab):.0f}%"],
            ["4",  "Internal Walls, Internal Plaster, Flooring",
             f"{get_cat_avg(brickwk):.0f}%"],
            ["5",  "Doors and Windows",
             f"{(get_cat_avg(doors) + get_cat_avg(windows)) / 2:.0f}%"],
            ["6",  "Sanitary and Electrical Fittings",
             f"{(get_cat_avg(plumbing) + get_cat_avg(elec)) / 2:.0f}%"],
            ["7",  "Tiling / Flooring",
             f"{get_cat_avg(tiling):.0f}%"],
            ["8",  "External Plumbing / Plaster, Elevation, Terrace Waterproofing, Painting",
             f"{(get_cat_avg(waterp) + get_cat_avg(painting)) / 2:.0f}%"],
            ["9",  "Carpark",
             f"{get_cat_avg(carpark):.0f}%"],
        ]
        add_data_table(doc, rows_a)

        doc.add_page_break()

    # TABLE B — Infrastructure
    add_section_heading(doc, "TABLE B: Internal & External Development Works in Respect of the Registered Phase")
    add_para(doc, "")

    infra = infrastructure_progress.get("activities", {}) if infrastructure_progress else {}

    def ic(key):
        return infra.get(key, {}).get("completion", 0)

    rows_b = [
        ["Sr. No.", "Common Areas and Facilities / Amenities", "Proposed\n(Yes/No)", "% Work Done", "Details"],
        ["1.",  "Internal Roads & Footpaths",               "Yes", f"{ic('road_footpath_storm_drain'):.0f}%",       ""],
        ["2.",  "Water Supply",                              "Yes", f"{ic('underground_water_distribution'):.0f}%",  ""],
        ["3.",  "Sewerage (STP)",                            "Yes", f"{ic('underground_sewage_network'):.0f}%",      ""],
        ["4.",  "Storm Water Drains",                        "Yes", f"{ic('road_footpath_storm_drain'):.0f}%",       ""],
        ["5.",  "Landscaping & Tree Planting",               "Yes", f"{ic('gardens_playground'):.0f}%",             ""],
        ["6.",  "Street Lighting",                           "Yes", f"{ic('street_lights'):.0f}%",                  ""],
        ["7.",  "Community Buildings (Club House)",          "Yes", f"{ic('club_house'):.0f}%",                     ""],
        ["8.",  "Treatment & Disposal of Sewage",            "Yes", f"{ic('sewage_treatment_plant'):.0f}%",         ""],
        ["9.",  "Solid Waste Management",                    "Yes", "0%",                                            ""],
        ["10.", "Rain Water Harvesting",                     "Yes", f"{ic('overhead_sump_reservoir'):.0f}%",        ""],
        ["11.", "Energy Management",                         "Yes", f"{ic('electric_substation_cables'):.0f}%",     ""],
        ["12.", "Fire Protection & Safety",                  "Yes", "0%",                                            ""],
        ["13.", "Electrical Meter Room / Substation",        "Yes", f"{ic('electric_substation_cables'):.0f}%",     ""],
        ["14.", "Swimming Pool",                             "Yes", f"{ic('swimming_pool'):.0f}%",                  ""],
        ["15.", "Amphitheatre",                              "Yes", f"{ic('amphitheatre'):.0f}%",                   ""],
        ["16.", "Boundary Wall & Entry Gate",                "Yes",
         f"{(ic('boundary_wall') + ic('entry_gate')) / 2:.0f}%",                                                    ""],
    ]
    add_data_table(doc, rows_b)

    add_para(doc, "")
    add_para(doc, "Yours Faithfully,", bold=True)
    add_para(doc, "")
    add_para(doc, "")
    add_para(doc, "_______________________________")
    add_para(doc, f"Signature & Name: {project.get('architect_name', '________________________')}", bold=True)
    add_para(doc, f"License No.: {project.get('architect_license', '________________________')}")

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


# ─── FORM 3 ──────────────────────────────────────────────────────────────────

def generate_form3_docx(project, buildings, construction_progress, infrastructure_progress, estimated_dev_cost, quarter, year):
    """Form-3: Cost Incurred = (% Work Completed from Construction Progress) × Estimated Cost"""
    doc = Document()
    section_margins(doc)

    add_title_block(doc, "FORM 3", "ENGINEER'S CERTIFICATE", "(Cost Incurred Statement for Development)")

    add_para(doc, "To", bold=True)
    for line in [
        project.get("promoter_name", "________________________"),
        project.get("promoter_address", project.get("address", "________________________")),
        f"{project.get('village', '')}, {project.get('taluka', '')}, {project.get('district', 'Goa')}",
    ]:
        add_para(doc, line, size=10)

    p_date = doc.add_paragraph()
    p_date.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p_date.add_run(f"Date: {datetime.now().strftime('%d/%m/%Y')}")
    r.font.bold = True
    r.font.size = Pt(10)

    add_para(doc, "")
    add_info_table(doc, [
        ("Project Name",            project.get("project_name", "")),
        ("RERA No.",                project.get("rera_number", "")),
        ("Report Period",           f"{quarter} {year}"),
        ("Report Date",             datetime.now().strftime("%d/%m/%Y")),
        ("Engineer Name",           project.get("engineer_name", "")),
        ("Engineer License No.",    project.get("engineer_license", "")),
    ])

    add_para(doc, "")
    add_para(doc, (
        f"I/We {project.get('engineer_name', '________________________')}, Engineer, have examined the "
        f"project site and hereby certify the cost incurred for development of the project "
        f"{project.get('project_name', '________')}, RERA No. {project.get('rera_number', '________')}."
    ), size=10)

    # TABLE A — Cost Incurred = completion% × estimated_cost
    add_para(doc, "")
    add_section_heading(doc, "TABLE A: Cost Incurred for Building Construction")

    cp_lookup = {cp.get("building_id"): cp for cp in (construction_progress or [])}
    total_est = total_inc = 0

    rows_a = [["Sr.", "Building / Wing", "% Complete", "Estimated Cost (₹)", "Cost Incurred (₹)", "Balance (₹)"]]
    for idx, b in enumerate(buildings or [], 1):
        progress = cp_lookup.get(b.get("building_id"), {})
        completion_pct = progress.get("overall_completion", 0)
        est = b.get("estimated_cost", 0)
        inc = round((completion_pct / 100) * est)
        bal = est - inc
        total_est += est
        total_inc += inc
        rows_a.append([str(idx), b.get("building_name", ""),
                        f"{completion_pct:.1f}%",
                        format_indian_number(est), format_indian_number(inc), format_indian_number(bal)])
    rows_a.append(["", "TOTAL", "",
                   format_indian_number(total_est), format_indian_number(total_inc),
                   format_indian_number(total_est - total_inc)])

    add_data_table(doc, rows_a)

    # TABLE B — Infra Cost Incurred = completion% × estimated_infra_cost
    add_para(doc, "")
    add_section_heading(doc, "TABLE B: Cost of Internal / External Development Works")

    est_c = estimated_dev_cost or {}
    infra_cost = est_c.get("infrastructure_cost", 0)
    infra_completion = (infrastructure_progress or {}).get("overall_completion", 0)
    infra_incurred = round((infra_completion / 100) * infra_cost)
    infra_balance = infra_cost - infra_incurred

    rows_b = [
        ["Sr.", "Development Work", "Estimated (₹)", "Incurred (₹)", "Balance (₹)"],
        ["1", "Internal Roads & Footpaths", "—", "—", "—"],
        ["2", "Water Supply & Distribution", "—", "—", "—"],
        ["3", "Sewerage & STP", "—", "—", "—"],
        ["4", "Storm Water Drains", "—", "—", "—"],
        ["5", "Landscaping", "—", "—", "—"],
        ["6", "Street Lighting", "—", "—", "—"],
        ["7", "Club House", "—", "—", "—"],
        ["8", "Swimming Pool", "—", "—", "—"],
        ["9", "Electrical Infrastructure", "—", "—", "—"],
        ["10", "Boundary Wall & Gate", "—", "—", "—"],
        ["", "TOTAL INFRASTRUCTURE",
         format_indian_number(infra_cost),
         format_indian_number(infra_incurred),
         format_indian_number(infra_balance)],
    ]
    add_data_table(doc, rows_b)

    add_para(doc, "")
    add_para(doc, "Yours Faithfully,", bold=True)
    add_para(doc, "")
    add_para(doc, "")
    add_para(doc, "_______________________________")
    add_para(doc, f"Signature & Name: {project.get('engineer_name', '________________________')}", bold=True)
    add_para(doc, f"License No.: {project.get('engineer_license', '________________________')}")

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


# ─── FORM 4 ──────────────────────────────────────────────────────────────────

def generate_form4_docx(project, form4_data, quarter, year):
    """
    FORM-4: CA Certificate in Word format.
    Matches the Excel/PDF CA Certificate format exactly:
    4-column table: Sr No / Particulars / Estimated Amount in Rs. / Incurred Amount in Rs.
    form4_data is a pre-computed dict from server._build_form4_data().
    """
    doc = Document()
    section_margins(doc)

    fd = form4_data or {}

    # Unpack all pre-computed values
    lc_a_est      = fd.get("lc_a_est", 0);    lc_a_inc      = fd.get("lc_a_inc", 0)
    lc_b_est      = fd.get("lc_b_est", 0);    lc_b_inc      = fd.get("lc_b_inc", 0)
    lc_c_est      = fd.get("lc_c_est", 0);    lc_c_inc      = fd.get("lc_c_inc", 0)
    lc_d_est      = fd.get("lc_d_est", 0);    lc_d_inc      = fd.get("lc_d_inc", 0)
    lc_e_est      = fd.get("lc_e_est", 0);    lc_e_inc      = fd.get("lc_e_inc", 0)
    rehab_i_est   = fd.get("rehab_i_est", 0); rehab_i_inc   = fd.get("rehab_i_inc", 0)
    rehab_ii_inc  = fd.get("rehab_ii_inc", 0)
    rehab_iii_inc = fd.get("rehab_iii_inc", 0)
    rehab_iv_inc  = fd.get("rehab_iv_inc", 0)
    rehab_any     = fd.get("rehab_any", False)
    land_sub_est  = fd.get("land_sub_est", 0); land_sub_inc  = fd.get("land_sub_inc", 0)
    dev_a1_est    = fd.get("dev_a1_est", 0)
    dev_a2_inc    = fd.get("dev_a2_inc", 0)
    dev_a3_est    = fd.get("dev_a3_est", 0);   dev_a3_inc    = fd.get("dev_a3_inc", 0)
    dev_b_est     = fd.get("dev_b_est", 0);    dev_b_inc     = fd.get("dev_b_inc", 0)
    dev_c_est     = fd.get("dev_c_est", 0);    dev_c_inc     = fd.get("dev_c_inc", 0)
    dev_sub_est   = fd.get("dev_sub_est", 0);  dev_sub_inc   = fd.get("dev_sub_inc", 0)
    total_est     = fd.get("total_est", 0);    total_inc     = fd.get("total_inc", 0)
    arch_pct      = fd.get("arch_pct", 0)
    proportion    = fd.get("proportion", 0)
    withdraw_allow = fd.get("withdraw_allow", 0)
    total_amount_received_sold = fd.get("total_amount_received_sold", 0)
    net_withdraw  = fd.get("net_withdraw", 0)
    bal_cost         = fd.get("bal_cost", 0)
    bal_recv_sold    = fd.get("bal_recv_sold", 0)
    unsold_area      = fd.get("unsold_area", 0)
    asr_rate         = fd.get("asr_rate", 0)
    avg_sale_price   = fd.get("avg_sale_price", 0)
    total_sale_val_sold = fd.get("total_sale_val_sold", 0)
    unsold_val       = fd.get("unsold_val", 0)
    total_recv       = fd.get("total_recv", 0)
    deposit_pct      = fd.get("deposit_pct", 0.70)
    deposit_amt      = fd.get("deposit_amt", 0)
    sales            = fd.get("sales", [])
    buildings        = fd.get("buildings", [])
    building_map_b   = {b.get("building_id"): b.get("building_name", "") for b in (buildings or [])}

    # ── Colour helpers ────────────────────────────────────────────────────────
    HEX_TITLE   = "1F3864"
    HEX_SECTION = "D9E1F2"
    HEX_SUBTOT  = "FCE4D6"
    HEX_SUMMARY = "E2EFDA"
    HEX_NET     = "FFFF00"
    HEX_NOTE    = "FFFACD"
    HEX_WHITE   = "FFFFFF"

    def fv(val, dec=False):
        if val is None or val == "NA":
            return "NA"
        if isinstance(val, str):
            return val
        if not val:
            return "NA"
        if dec:
            return f"{float(val):,.2f}"
        return format_indian_number(int(val))

    def fpct(val):
        if not val:
            return "NA"
        return f"{float(val)*100:.2f}%"

    # ── Table creation helper ─────────────────────────────────────────────────
    # We build one large 4-column table for the whole Form-4

    def make_table(n_rows):
        t = doc.add_table(rows=n_rows, cols=4)
        t.style = "Table Grid"
        # Column widths: Sr / Particulars / Estimated / Incurred
        col_widths = [Cm(1.2), Cm(10.2), Cm(3.5), Cm(3.5)]
        for row in t.rows:
            row.cells[0].width = col_widths[0]
            for ci, cw in enumerate(col_widths):
                tc = row.cells[ci]._tc
                tcPr = tc.get_or_add_tcPr()
                from docx.oxml import OxmlElement as _OE
                from docx.oxml.ns import qn as _qn
                tcW = _OE("w:tcW")
                tcW.set(_qn("w:w"), str(int(cw.emu / 914.4 * 20)))
                tcW.set(_qn("w:type"), "dxa")
                tcPr.append(tcW)
        return t

    def _write_cell(cell, text, bold=False, size=8, align=WD_ALIGN_PARAGRAPH.LEFT,
                    fg=None, bg=None, italic=False):
        cell.paragraphs[0].clear()
        para = cell.paragraphs[0]
        para.alignment = align
        run = para.add_run(str(text) if text is not None else "")
        run.font.size = Pt(size)
        run.font.bold = bold
        run.font.italic = italic
        if fg:
            run.font.color.rgb = RGBColor(*bytes.fromhex(fg))
        if bg:
            shade_cell(cell, bg)
        border_cell(cell)

    def _span_row(table, row_idx, text, bold=False, size=9, align=WD_ALIGN_PARAGRAPH.CENTER,
                  bg=None, fg_hex=None, italic=False):
        """Merge all 4 cells in a row and write text."""
        row = table.rows[row_idx]
        row.cells[0].merge(row.cells[3])
        _write_cell(row.cells[0], text, bold=bold, size=size, align=align,
                    bg=bg, fg=fg_hex, italic=italic)

    def _data_row(table, row_idx, sr, particulars, est, inc,
                  bg=None, bold=False, size=8, note=False):
        row = table.rows[row_idx]
        _write_cell(row.cells[0], sr, bold=bold, size=size,
                    align=WD_ALIGN_PARAGRAPH.CENTER, bg=bg, italic=note)
        _write_cell(row.cells[1], particulars, bold=bold, size=size, bg=bg, italic=note)
        _write_cell(row.cells[2], est, bold=bold, size=size,
                    align=WD_ALIGN_PARAGRAPH.RIGHT, bg=bg, italic=note)
        _write_cell(row.cells[3], inc, bold=bold, size=size,
                    align=WD_ALIGN_PARAGRAPH.RIGHT, bg=bg, italic=note)

    # Pre-count rows needed
    # Title(1) + subtitle(1) + project info(1) + col headers(1)
    # Land section header(1) + 5 items + rehab header(1) + 4 rehab items + rehab note(2) + sub-total(1) = 16
    # Dev section header(1) + 2 items + note(1) + on-site(1) + taxes(1) + finance(1) + sub-total(1) = 8
    # Summary rows Sr2-8(7)
    # Sig block 1(6 lines) = 6
    # Addl info header(1) + 5 addl rows + 1 deposit row = 7
    # Sig block 2(6 lines) = 6
    # Annexure A header(2) + sold header(2) + up to 20 sold rows + total(1) + unsold box(1)
    # + unsold header(1) + up to 20 unsold rows + total(1) = variable
    sold_sales   = [s for s in (sales or []) if s.get("buyer_name")]
    unsold_sales = [s for s in (sales or []) if not s.get("buyer_name")]
    n_sold   = max(len(sold_sales), 3)
    n_unsold = max(len(unsold_sales), 3)

    FIXED_ROWS = (4 + 16 + 8 + 7 + 6 + 7 + 6)
    ANNEX_ROWS = 2 + 2 + n_sold + 1 + 1 + 1 + n_unsold + 1
    total_rows  = FIXED_ROWS + ANNEX_ROWS

    t = make_table(total_rows)
    ri = 0  # row index pointer

    # ── Title ─────────────────────────────────────────────────────
    _span_row(t, ri, "Form-4: CA Certificate", bold=True, size=12, bg=HEX_TITLE, fg_hex=HEX_WHITE); ri += 1
    _span_row(t, ri, "(FOR REGISTRATION OF A PROJECT AND SUBSEQUENT WITHDRAWAL OF MONEY)",
              bold=True, size=9, bg=HEX_TITLE, fg_hex=HEX_WHITE); ri += 1

    proj_info = (f"Project: {project.get('project_name','—')}   |   RERA No.: {project.get('rera_number','—')}"
                 f"   |   Period: {quarter} {year}   |   Date: {datetime.now().strftime('%d/%m/%Y')}")
    _span_row(t, ri, proj_info, size=8, bg=HEX_SECTION); ri += 1

    # Col headers
    _data_row(t, ri, "Sr No", "Particulars", "Estimated Amount in Rs.", "Incurred Amount in Rs.",
              bg=HEX_SECTION, bold=True, size=9); ri += 1

    # ── Land Cost Section ─────────────────────────────────────────
    row = t.rows[ri]
    row.cells[0].merge(row.cells[0])  # keep Sr col
    _write_cell(row.cells[0], "", bg=HEX_SECTION)
    row.cells[1].merge(row.cells[3])
    _write_cell(row.cells[1], "i.  Land Cost :", bold=True, size=9, bg=HEX_SECTION); ri += 1

    def item(sr, part, est_v, inc_v, bg=None, note=False):
        nonlocal ri
        _data_row(t, ri, sr, part, fv(est_v), fv(inc_v, dec=True), bg=bg, note=note); ri += 1

    item(1,
         "a. Acquisition Cost of Land or Development Rights, lease Premium, lease rent, "
         "interest cost incurred or payable on Land Cost and legal cost.",
         lc_a_est, lc_a_inc)

    item("",
         "b. Amount of Premium payable to obtain development rights, FSI, additional FSI, "
         "fungible area, and any other incentive under DCR from Local Authority or State "
         "Government or any Statutory Authority.",
         lc_b_est, lc_b_inc)

    item("", "c. Acquisition cost of TDR (if any)", lc_c_est, lc_c_inc)

    item("",
         "d. Amounts payable to State Government or competent authority or any other statutory "
         "authority of the State or Central Government, towards stamp duty, transfer charges, "
         "registration fees etc; and",
         lc_d_est, lc_d_inc)

    item("",
         "e. Land Premium payable as per annual statement of rates (ASR) for redevelopment "
         "of land owned by public authorities.",
         lc_e_est, lc_e_inc)

    # Rehab header
    _data_row(t, ri, "", "f. Under Rehabilitation Scheme :", "", "", bold=True, size=8); ri += 1

    item("",
         "   (i) Estimated construction cost of rehab building including site development "
         "and infrastructure for the same as certified by Engineer.",
         rehab_i_est if rehab_any else "NA",
         rehab_i_inc if rehab_any else "NA")

    item("",
         "   (ii) Actual Cost of construction of rehab building incurred as per the books "
         "of accounts as verified by the CA.",
         "NA", rehab_ii_inc if rehab_any else "NA")

    # Rehab note
    row = t.rows[ri]; row.cells[1].merge(row.cells[3])
    _write_cell(row.cells[0], "", bg=HEX_NOTE)
    _write_cell(row.cells[1],
                "Note: (for total cost of construction incurred, Minimum of (i) or (ii) is to be considered).",
                size=7, italic=True, bg=HEX_NOTE); ri += 1

    item("",
         "   (iii) Cost towards clearance of land of all or any encumbrances including cost of "
         "removal of legal/illegal occupants, cost for providing temporary transit accommodation "
         "or rent in lieu of Transit Accommodation, overhead cost,",
         "NA", rehab_iii_inc if rehab_any else "NA")

    item("",
         "   (iv) Cost of ASR linked premium, fees, charges and security deposits or maintenance "
         "deposit, or any amount whatsoever payable to any authorities towards and in project of "
         "rehabilitation.",
         "NA", rehab_iv_inc if rehab_any else "NA")

    # Second rehab note (blank / placeholder)
    row = t.rows[ri]; row.cells[1].merge(row.cells[3])
    _write_cell(row.cells[0], "", bg=HEX_NOTE)
    _write_cell(row.cells[1], "", bg=HEX_NOTE); ri += 1

    # Land sub-total
    _data_row(t, ri, "", "Sub-Total of LAND COST",
              fv(land_sub_est), fv(land_sub_inc, dec=True),
              bg=HEX_SUBTOT, bold=True, size=9); ri += 1

    # ── Development Cost Section ──────────────────────────────────
    row = t.rows[ri]
    _write_cell(row.cells[0], "", bg=HEX_SECTION)
    row.cells[1].merge(row.cells[3])
    _write_cell(row.cells[1], "ii.  Development Cost / Cost of Construction :",
                bold=True, size=9, bg=HEX_SECTION); ri += 1

    item("",
         "a. (i) Estimated Cost of Construction as certified by Engineer.",
         dev_a1_est, "Refer Note")

    item("",
         "   (ii) Actual Cost of construction incurred as per the books of accounts "
         "as verified by the CA.",
         "Refer Note", dev_a2_inc)

    # MIN note
    row = t.rows[ri]; row.cells[1].merge(row.cells[3])
    _write_cell(row.cells[0], "", bg=HEX_NOTE)
    _write_cell(row.cells[1],
                "Note: (for adding to total cost of construction incurred, Minimum of (i) or (ii) is to be considered).",
                size=7, italic=True, bg=HEX_NOTE); ri += 1

    item("",
         "(iii) On-site expenditure for development of entire project excluding cost of "
         "construction as per (i) or (ii) above, i.e. salaries, consultants fees, site "
         "overheads, development works, cost of services (including water, electricity, "
         "sewerage, drainage, layout roads etc.), cost of machineries and equipment "
         "including its hire and maintenance costs, consumables etc.",
         dev_a3_est, dev_a3_inc)

    item("",
         "b. Payment of Taxes, cess, fees, charges, premiums, interest etc. "
         "to any statutory Authority.",
         dev_b_est, dev_b_inc)

    item("",
         "c. Principal sum and interest payable to financial institutions, scheduled banks, "
         "non-banking financial institution (NBFC) or money lenders on construction funding "
         "or money borrowed for construction;",
         dev_c_est, dev_c_inc)

    _data_row(t, ri, "", "Sub-Total of Development Cost",
              fv(dev_sub_est), fv(dev_sub_inc, dec=True),
              bg=HEX_SUBTOT, bold=True, size=9); ri += 1

    # ── Summary Rows Sr 2-8 ───────────────────────────────────────
    def summ(sr, text, est_v="", inc_v="", net=False):
        nonlocal ri
        bg = HEX_NET if net else HEX_SUMMARY
        _data_row(t, ri, str(sr), text, est_v, inc_v, bg=bg, bold=net, size=8); ri += 1

    summ(2, "Total Estimated Cost of the Real Estate Project [1(i) + 1(ii)] of Estimated Column.",
         fv(total_est))
    summ(3, "Total Cost Incurred of the Real Estate Project [1(i) + 1(ii)] of Incurred Column.",
         inc_v=fv(total_inc, dec=True))
    summ(4, "% Completion of Construction Work (as per Project Architect's Certificate)",
         inc_v=fpct(arch_pct))
    summ(5,
         f"Proportion of the Cost incurred on Land Cost and {arch_pct*100:.2f}% "
         f"Construction Cost to the Total Estimated Cost.  (Sr.3 / Sr.2 %)",
         inc_v=fpct(proportion))
    summ(6, "Amount Which can be Withdrawn from the Designated Account. "
         "(Total Estimated Cost × Proportion of cost incurred  =  Sr.2 × Sr.5)",
         inc_v=fv(withdraw_allow, dec=True))
    summ(7, "Less: Total Sale Amount Received from Sold Units "
         "(as per Annexure A – sum of amounts received from all sold unit allottees).",
         inc_v=fv(total_amount_received_sold, dec=True))
    summ(8, "Net Amount which can be Withdrawn from the Designated Bank Account "
         "under this Certificate.  (Sr.6 – Sr.7)",
         inc_v=fv(net_withdraw, dec=True), net=True)

    # ── Signature block 1 ─────────────────────────────────────────
    cert_text = (f"This certificate is being issued for RERA compliance for the Company "
                 f"[{project.get('promoter_name', 'Promoter')}] and is based on the records "
                 f"and documents produced before me and explanations provided to me by the "
                 f"management of the Company.")
    _span_row(t, ri, cert_text, size=8, align=WD_ALIGN_PARAGRAPH.LEFT); ri += 1

    for sig_line in ["Yours Faithfully,",
                     f"Signature of Chartered Accountant – {project.get('ca_name', '')}",
                     f"(Membership No.: {project.get('ca_membership', '…………')})",
                     "______________________", "Name"]:
        row = t.rows[ri]
        _write_cell(row.cells[0], ""); _write_cell(row.cells[1], "")
        row.cells[2].merge(row.cells[3])
        _write_cell(row.cells[2], sig_line, size=8,
                    bold=(sig_line in ["Yours Faithfully,", "Name"]),
                    align=WD_ALIGN_PARAGRAPH.CENTER); ri += 1

    # ── Additional Information ────────────────────────────────────
    _span_row(t, ri, "(ADDITIONAL INFORMATION FOR ONGOING PROJECTS)",
              bold=True, size=10, bg=HEX_TITLE, fg_hex=HEX_WHITE); ri += 1

    def addl(sr, text, est_v="", inc_v=""):
        nonlocal ri
        _data_row(t, ri, str(sr) if sr is not None else "", text, est_v, inc_v, size=8); ri += 1

    addl(1, "Estimated Balance Cost to Complete the Real Estate Project "
         "(Difference of Total Estimated Project cost less Cost incurred) (calculated as per Form IV)",
         inc_v=fv(bal_cost, dec=True))
    addl(2, "Balance amount of receivables from sold apartments as per Annexure A to this certificate "
         "(as certified by Chartered Accountant as verified from the records and books of Accounts)",
         inc_v=fv(bal_recv_sold, dec=True) if bal_recv_sold else "NIL")
    addl(3, "(i) Balance Unsold area (to be certified by Management and verified by CA)",
         inc_v=f"{unsold_area:,.2f} sq.m." if unsold_area else "NIL")
    addl(None,
         f"(ii) Estimated amount of sales proceeds in respect of unsold apartments "
         f"(Avg Sale Price per sq.m. of Sold Units: "
         f"Rs.{format_indian_number(int(avg_sale_price)) if avg_sale_price else '—'} × "
         f"Unsold Area: {unsold_area:,.2f} sq.m.) as per Annexure A",
         inc_v=fv(unsold_val, dec=True) if unsold_val else "NIL")
    addl(4,
         f"Estimated receivables of ongoing project.  "
         f"Total Sale Value of Sold Units (Rs.{format_indian_number(int(total_sale_val_sold)) if total_sale_val_sold else '0'}) + 3(ii)",
         inc_v=fv(total_recv, dec=True) if total_recv else "NIL")
    addl(5, "Amount to be deposited in Designated Account – 70% or 100%")

    if total_recv > bal_cost:
        dep_text = (f"IF Sr.4 is GREATER THAN Sr.1: 70% × Rs.{format_indian_number(int(total_recv))} "
                    f"= Rs.{format_indian_number(int(total_recv * 0.70))}")
    else:
        dep_text = (f"IF Sr.4 is LESSER THAN or EQUAL TO Sr.1: 100% × Rs.{format_indian_number(int(total_recv))} "
                    f"= Rs.{format_indian_number(int(total_recv))}")
    _span_row(t, ri, dep_text, bold=True, size=8, align=WD_ALIGN_PARAGRAPH.LEFT); ri += 1

    # ── Signature block 2 ─────────────────────────────────────────
    _span_row(t, ri, cert_text, size=8, align=WD_ALIGN_PARAGRAPH.LEFT); ri += 1

    for sig_line in ["Yours Faithfully,",
                     "Signature of Chartered Accountant,",
                     f"(Membership No.: {project.get('ca_membership', '…………')})",
                     "______________________", "Name"]:
        row = t.rows[ri]
        _write_cell(row.cells[0], ""); _write_cell(row.cells[1], "")
        row.cells[2].merge(row.cells[3])
        _write_cell(row.cells[2], sig_line, size=8,
                    bold=(sig_line in ["Yours Faithfully,", "Name"]),
                    align=WD_ALIGN_PARAGRAPH.CENTER); ri += 1

    # ── Annexure A ────────────────────────────────────────────────
    _span_row(t, ri, "Annexure A", bold=True, size=12, bg=HEX_TITLE, fg_hex=HEX_WHITE); ri += 1
    _span_row(t, ri,
              "Statement for calculation of Receivables from the Sales of the Ongoing Real Estate Project",
              bold=True, size=9); ri += 1
    _span_row(t, ri, "Sold Inventory", bold=True, size=9, bg=HEX_SECTION); ri += 1

    # Annexure A column headers
    _data_row(t, ri, "Sr No", "Flat No.", "Received Amount (Rs.)", "Balance Receivable (Rs.)",
              bg=HEX_SECTION, bold=True, size=8); ri += 1

    total_recv_amt = 0; total_bal_recv = 0
    if sold_sales:
        for idx, s in enumerate(sold_sales, 1):
            bname = building_map_b.get(s.get("building_id"), s.get("building_name", ""))
            unit  = f"{bname} – {s.get('unit_number','')}" if bname else s.get("unit_number", "")
            recv  = s.get("amount_received", 0)
            bal   = s.get("sale_value", 0) - recv
            total_recv_amt += recv; total_bal_recv += bal
            _data_row(t, ri, str(idx), unit, fv(recv, dec=True), fv(bal, dec=True), size=8); ri += 1
    else:
        for i in range(1, n_sold + 1):
            _data_row(t, ri, str(i), "", "", "", size=8); ri += 1

    _data_row(t, ri, "", "Total", fv(total_recv_amt, dec=True), fv(total_bal_recv, dec=True),
              bg=HEX_SUBTOT, bold=True, size=8); ri += 1

    # Unsold box
    unsold_box = (f"(Unsold Inventory Valuation)\n"
                  f"Average Sale Price per sq.m. of Sold Units: "
                  f"Rs. {format_indian_number(int(avg_sale_price)) if avg_sale_price else '—'} per sq.m.\n"
                  f"Total Unsold Area: {unsold_area:,.2f} sq.m.\n"
                  f"Estimated Unsold Inventory Value (Avg Price × Unsold Area): "
                  f"Rs. {format_indian_number(int(unsold_val)) if unsold_val else 'NIL'}")
    _span_row(t, ri, unsold_box, size=8, align=WD_ALIGN_PARAGRAPH.LEFT); ri += 1

    _data_row(t, ri, "Sr No", "Flat Number", "Area (sq.m.)", "Estimated Value (Rs.)",
              bg=HEX_SECTION, bold=True, size=8); ri += 1

    total_u_area = 0; total_u_val = 0
    if unsold_sales:
        for idx, s in enumerate(unsold_sales, 1):
            area = s.get("carpet_area", 0) or 0
            uval = area * avg_sale_price if avg_sale_price else 0
            total_u_area += area; total_u_val += uval
            bname = building_map_b.get(s.get("building_id"), s.get("building_name", ""))
            unit  = f"{bname} – {s.get('unit_number','')}" if bname else s.get("unit_number", "")
            _data_row(t, ri, str(idx), unit, f"{area:,.2f}" if area else "",
                      fv(uval, dec=True) if uval else "", size=8); ri += 1
    else:
        for i in range(1, n_unsold + 1):
            _data_row(t, ri, str(i), "", "", "", size=8); ri += 1

    _data_row(t, ri, "", "Total",
              f"{total_u_area:,.2f}" if total_u_area else "",
              fv(total_u_val, dec=True) if total_u_val else "",
              bg=HEX_SUBTOT, bold=True, size=8); ri += 1

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


# ─── ANNEXURE-A ──────────────────────────────────────────────────────────────

def generate_annexure_a_docx(project, sales, buildings, quarter, year):
    doc = Document()

    # Landscape layout
    for section in doc.sections:
        section.top_margin    = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin   = Cm(1.5)
        section.right_margin  = Cm(1.5)
        # Swap width/height for landscape
        w = section.page_width
        h = section.page_height
        section.page_width  = max(w, h)
        section.page_height = min(w, h)

    add_para(doc, "ANNEXURE - A", size=14, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc, "Statement of Receivables from Allottees", size=12, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc, (
        f"Project: {project.get('project_name', '')}  |  "
        f"RERA No: {project.get('rera_number', '')}  |  "
        f"Period: {quarter} {year}  |  Date: {datetime.now().strftime('%d/%m/%Y')}"
    ), size=10)
    add_para(doc, "")

    headers = [
        "Sr.", "Unit No.", "Building / Wing", "Type", "Area (sq.m.)",
        "Agreement Value (₹)", "Amount Received (₹)", "Balance Due (₹)",
        "Due Date", "Allottee Name",
    ]
    bl = {b.get("building_id"): b.get("building_name", "") for b in (buildings or [])}

    total_val = total_recv = total_bal = 0
    all_rows = [headers]

    for idx, sale in enumerate(sales or [], 1):
        agr  = sale.get("sale_value", sale.get("agreement_value", 0))
        recv = sale.get("amount_received", 0)
        bal  = agr - recv
        total_val  += agr
        total_recv += recv
        total_bal  += bal
        all_rows.append([
            str(idx),
            sale.get("unit_number", ""),
            bl.get(sale.get("building_id"), sale.get("building_name", "")),
            sale.get("unit_type", sale.get("flat_type", "")),
            str(sale.get("carpet_area", "")),
            format_indian_number(agr),
            format_indian_number(recv),
            format_indian_number(bal),
            sale.get("agreement_date", sale.get("due_date", "")),
            sale.get("buyer_name", sale.get("allottee_name", "")),
        ])

    if not sales:
        all_rows.append(["", "No sales data available"] + [""] * 8)

    all_rows.append([
        "", "", "", "", "TOTAL",
        format_indian_number(total_val),
        format_indian_number(total_recv),
        format_indian_number(total_bal),
        "", "",
    ])

    add_data_table(doc, all_rows)

    add_para(doc, "")
    add_para(doc, (
        f"Summary: Total Units: {len(sales or [])}  |  "
        f"Total Agreement Value: ₹{format_indian_number(total_val)}  |  "
        f"Total Received: ₹{format_indian_number(total_recv)}  |  "
        f"Total Balance Due: ₹{format_indian_number(total_bal)}"
    ), size=10, bold=True)

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf
